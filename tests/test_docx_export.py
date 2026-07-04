import io
import unittest

from docx import Document

from estate_planning.docx_export import render_docx
from estate_planning.documents import MODE_EN, generate
from estate_planning.models import Asset, Beneficiary, EstatePlan


def sample():
    return EstatePlan(
        full_name="Jane",
        nationality="USA",
        passport_or_id_number="P1",
        date_of_birth="1980-01-01",
        thai_address="BKK",
        status="foreign_resident",
        beneficiaries=[Beneficiary("Somchai", "descendant")],
        assets=[Asset("bank_account", "BBL", 850000, beneficiary="Somchai")],
    )


class DocxExportTests(unittest.TestCase):
    def test_produces_valid_docx(self):
        _, md = generate(sample(), ["last_will"], MODE_EN)["last_will"]
        data = render_docx(md)
        self.assertEqual(data[:2], b"PK")  # zip signature
        doc = Document(io.BytesIO(data))
        self.assertTrue(doc.paragraphs)

    def test_headings_become_word_headings(self):
        _, md = generate(sample(), ["last_will"], MODE_EN)["last_will"]
        doc = Document(io.BytesIO(render_docx(md)))
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        self.assertTrue(headings)

    def test_asset_inventory_table_rendered(self):
        _, md = generate(sample(), ["asset_inventory"], MODE_EN)["asset_inventory"]
        doc = Document(io.BytesIO(render_docx(md)))
        self.assertEqual(len(doc.tables), 1)
        header_cells = [c.text for c in doc.tables[0].rows[0].cells]
        self.assertIn("Category", header_cells)

    def test_bold_inline_preserved(self):
        doc = Document(io.BytesIO(render_docx("This is **bold** text.")))
        runs = doc.paragraphs[0].runs
        self.assertTrue(any(r.bold and r.text == "bold" for r in runs))

    def test_all_documents_export(self):
        docs = generate(sample(), None, MODE_EN)
        for key, (_title, md) in docs.items():
            data = render_docx(md)
            self.assertEqual(data[:2], b"PK", f"{key} not a valid docx")


if __name__ == "__main__":
    unittest.main()
