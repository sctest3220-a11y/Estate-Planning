"""Convert a generated document's Markdown into a Word (.docx) file.

Handles the subset of Markdown produced by documents.py: # / ## headings,
**bold** / *italic* inline, "|" tables, "> " blockquotes, "- " bullets (with one
level of nesting), "---" separators, and plain paragraphs.
"""

import io
import re

from docx import Document
from docx.shared import Inches

_INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def _add_inline(paragraph, text):
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def _is_separator_row(cells):
    return all(set(c) <= set("-: ") and c for c in cells)


def _add_table(doc, table_lines):
    rows = [[c.strip() for c in tl.strip().strip("|").split("|")] for tl in table_lines]
    rows = [r for r in rows if not _is_separator_row(r)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    for i, r in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncols):
            value = r[j] if j < len(r) else ""
            para = cells[j].paragraphs[0]
            _add_inline(para, value)
            if i == 0:
                for run in para.runs:
                    run.bold = True


def render_docx(md_text):
    """Return a single document's Markdown rendered as .docx bytes."""
    doc = Document()
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            _add_inline(p, stripped[2:])
            for run in p.runs:
                run.italic = True
        elif stripped.startswith("- "):
            indent = len(raw) - len(raw.lstrip())
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped[2:])
        elif stripped == "---":
            pass  # section separator — skip
        else:
            _add_inline(doc.add_paragraph(), stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
